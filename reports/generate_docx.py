"""
Generate Word report for AP Transco DGA Fleet Health Assessment.
Professional multi-section document with risk reasons.

Usage: python reports/generate_docx.py
Output: reports/AP_Transco_DGA_Fleet_Health_Report.docx
"""
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_helpers import load_data, build_risk_reason, build_short_reason, REPORTS_DIR

TODAY = datetime.now().strftime("%d %B %Y")
BLUE = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xE6, 0x51, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
RISK_COLORS = {"Excellent": GREEN, "Good": RGBColor(0x15, 0x65, 0xC0), "Fair": AMBER, "Poor": RED, "Critical": RGBColor(0xB7, 0x1C, 0x1C)}
RISK_ORDER = {"Critical": 0, "Poor": 1, "Fair": 2, "Good": 3, "Excellent": 4}


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex
    })
    shading.append(shd)


def add_table_row(table, values, bold=False, header=False, risk_col=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(val) if val is not None else "-")
        run.font.size = Pt(8)
        run.font.name = "Arial"

        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "1F4E79")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif bold:
            run.bold = True

        if risk_col is not None and i == risk_col:
            risk_val = str(val)
            if risk_val in RISK_COLORS:
                run.font.color.rgb = RISK_COLORS[risk_val]
                run.bold = True
    return row


def main():
    data = load_data()
    fleet = data["fleet"]
    transformers = data["transformers"]
    substations = data["substations"]
    models = data["models"]

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ANDHRA PRADESH TRANSMISSION CORPORATION")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(AP TRANSCO)")
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Transformer Fleet Health Assessment Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dissolved Gas Analysis (DGA) Based Condition Assessment")
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Using Machine Learning & IEEE/IEC Standards")
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Report Date: {TODAY}")
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Data Period: {fleet.get('date_range',{}).get('min','N/A')} to {fleet.get('date_range',{}).get('max','N/A')}")
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL")
    r.bold = True
    r.font.color.rgb = RED

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    total = fleet.get("total_transformers", 0)
    samples = fleet.get("total_samples", 0)
    subs_count = fleet.get("total_substations", 0)
    risk_dist = fleet.get("risk_distribution", {})
    fault_dist = fleet.get("fault_distribution", {})
    poor_count = risk_dist.get("Poor", 0) + risk_dist.get("Critical", 0)

    doc.add_paragraph(
        f"This report presents the results of a comprehensive machine learning-based condition assessment "
        f"of AP Transco\u2019s power transformer fleet. The analysis covers {total:,} transformers across "
        f"{subs_count} substations, based on {samples:,} oil test samples."
    )

    doc.add_heading("Key Findings", level=2)
    findings = [
        f"{risk_dist.get('Excellent', 0):,} transformers ({risk_dist.get('Excellent', 0)/total*100:.0f}%) in Excellent health (CHI 80\u2013100)",
        f"{poor_count} transformers require attention (Poor or Critical health)",
        f"Average fleet Composite Health Index: {fleet.get('avg_chi', 0):.1f} / 100",
        f"Fault classification model: {models.get('fault_classifier',{}).get('accuracy',0)*100:.1f}% accuracy across 8 fault types",
        f"{len([s for s in substations if s.get('worst_risk')=='Poor'])} substations contain at least one Poor transformer",
        f"Primary risk factor: Elevated moisture content in transformer oil",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    # Risk distribution table
    doc.add_heading("Fleet Risk Distribution", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    add_table_row(table, ["Risk Level", "Count", "Percentage", "Action Required"], header=True)
    for level in ["Excellent", "Good", "Fair", "Poor", "Critical"]:
        count = risk_dist.get(level, 0)
        pct = f"{count/total*100:.1f}%"
        action = {"Excellent": "No action", "Good": "Continue monitoring", "Fair": "Increase test frequency",
                  "Poor": "Plan maintenance", "Critical": "Immediate action"}[level]
        add_table_row(table, [level, count, pct, action], risk_col=0)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 2. FAULT DISTRIBUTION
    # ════════════════════════════════════════════════════════
    doc.add_heading("2. Fault Type Distribution", level=1)
    doc.add_paragraph(
        "Each transformer\u2019s latest DGA sample is classified into one of 8 fault types per IEC 60599:2022."
    )

    fault_table = doc.add_table(rows=1, cols=4)
    fault_table.style = "Table Grid"
    add_table_row(fault_table, ["Fault Code", "Description", "Count", "Key Gases"], header=True)
    fault_info = [
        ("Normal", "No Fault Detected", "Below IEEE Status 1"),
        ("T1", "Thermal < 300\u00B0C", "CH4, C2H6"),
        ("T3", "Thermal > 700\u00B0C", "C2H4, CH4"),
        ("D2", "High-energy Discharge", "H2, C2H2, C2H4"),
        ("D1", "Low-energy Discharge", "H2, C2H2"),
        ("PD", "Partial Discharge", "H2, CH4"),
        ("T2", "Thermal 300\u2013700\u00B0C", "CH4, C2H4, C2H6"),
        ("DT", "Discharge + Thermal", "C2H2, C2H4, CH4"),
    ]
    for code, desc, gases in fault_info:
        count = fault_dist.get(code, 0)
        add_table_row(fault_table, [code, desc, count, gases])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 3. TRANSFORMERS REQUIRING ATTENTION (WITH REASON)
    # ════════════════════════════════════════════════════════
    doc.add_heading("3. Transformers Requiring Attention", level=1)

    worst = sorted(
        [t for t in transformers if t.get("risk_level") in ("Poor", "Critical")],
        key=lambda t: t.get("chi") or 999
    )

    doc.add_paragraph(
        f"The following {len(worst)} transformers are rated Poor or Critical. "
        f"The \u2018Reason for Risk Rating\u2019 column explains which parameter exceeds IEEE/IEC thresholds."
    )

    att_table = doc.add_table(rows=1, cols=7)
    att_table.style = "Table Grid"
    add_table_row(att_table, ["Equipment", "Substation", "Make", "kV", "CHI", "Risk", "Reason for Risk Rating"], header=True)

    # Set column widths
    for row in att_table.rows:
        row.cells[0].width = Inches(0.9)
        row.cells[1].width = Inches(1.8)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(0.4)
        row.cells[4].width = Inches(0.4)
        row.cells[5].width = Inches(0.5)
        row.cells[6].width = Inches(1.8)

    for t in worst:
        reason = build_risk_reason(t)
        r = add_table_row(att_table, [
            t.get("equipment_no"),
            (t.get("substation_name") or t.get("substation_id") or "-")[:35],
            t.get("make") or "-",
            t.get("voltage_class") or "-",
            f"{t['chi']:.1f}" if t.get("chi") is not None else "-",
            t.get("risk_level") or "-",
            reason,
        ], risk_col=5)
        # Color the reason cell red
        reason_run = r.cells[6].paragraphs[0].runs[0]
        reason_run.font.color.rgb = RED
        reason_run.font.size = Pt(7)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 4. SUBSTATIONS WITH POOR RISK
    # ════════════════════════════════════════════════════════
    doc.add_heading("4. Substations with Poor Transformers", level=1)
    poor_subs = [s for s in substations if s.get("worst_risk") == "Poor"]
    doc.add_paragraph(f"{len(poor_subs)} substations contain at least one transformer rated Poor.")

    sub_table = doc.add_table(rows=1, cols=6)
    sub_table.style = "Table Grid"
    add_table_row(sub_table, ["Substation", "kV", "Units", "Avg CHI", "Worst", "Last Tested"], header=True)
    for s in poor_subs:
        add_table_row(sub_table, [
            (s.get("name") or s.get("id"))[:40],
            s.get("voltage_class") or "-",
            s.get("transformer_count"),
            f"{s['avg_chi']:.1f}" if s.get("avg_chi") else "-",
            s.get("worst_risk"),
            s.get("latest_sample_date") or "-",
        ], risk_col=4)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 5. MANUFACTURER ANALYSIS
    # ════════════════════════════════════════════════════════
    doc.add_heading("5. Manufacturer Performance", level=1)

    from collections import defaultdict
    make_map = defaultdict(lambda: {"n": 0, "chi_s": 0, "chi_n": 0, "faults": 0, "poor": 0})
    for t in transformers:
        m = t.get("make") or "Unknown"
        make_map[m]["n"] += 1
        if t.get("chi") is not None: make_map[m]["chi_s"] += t["chi"]; make_map[m]["chi_n"] += 1
        if t.get("fault_label") != "Normal": make_map[m]["faults"] += 1
        if t.get("risk_level") in ("Poor", "Critical"): make_map[m]["poor"] += 1

    make_table = doc.add_table(rows=1, cols=6)
    make_table.style = "Table Grid"
    add_table_row(make_table, ["Manufacturer", "Units", "Avg CHI", "Fault Rate", "Poor/Crit", "Assessment"], header=True)

    for make, d in sorted(make_map.items(), key=lambda x: -x[1]["n"]):
        if d["n"] < 10:
            continue
        avg_chi = d["chi_s"] / d["chi_n"] if d["chi_n"] else 0
        fr = d["faults"] / d["n"] * 100
        assess = "Needs review" if d["poor"] > 0 else "High fault rate" if fr > 40 else "Below avg" if avg_chi < 85 else "OK"
        add_table_row(make_table, [make, d["n"], f"{avg_chi:.1f}", f"{fr:.1f}%", d["poor"], assess])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 6. MODEL PERFORMANCE
    # ════════════════════════════════════════════════════════
    doc.add_heading("6. Machine Learning Model Performance", level=1)

    fc = models.get("fault_classifier", {})
    hi = models.get("health_index", {})
    fp = models.get("failure_predictor", {})

    doc.add_heading("Model 1: Fault Classification", level=2)
    doc.add_paragraph(f"Accuracy: {fc.get('accuracy',0)*100:.1f}% | Weighted F1: {fc.get('weighted_f1',0)*100:.1f}% | Macro F1: {fc.get('macro_f1',0)*100:.1f}%")

    fc_table = doc.add_table(rows=1, cols=5)
    fc_table.style = "Table Grid"
    add_table_row(fc_table, ["Fault", "Precision", "Recall", "F1", "Support"], header=True)
    for cls, m in (fc.get("per_class") or {}).items():
        add_table_row(fc_table, [cls, f"{m['precision']*100:.1f}%", f"{m['recall']*100:.1f}%",
                                  f"{m['f1-score']*100:.1f}%", int(m["support"])])

    doc.add_heading("Model 2: Health Index Regression", level=2)
    doc.add_paragraph(f"R\u00B2: {hi.get('r2',0):.4f} | RMSE: {hi.get('rmse',0):.2f} | MAE: {hi.get('mae',0):.2f} | Risk Accuracy: {hi.get('risk_level_accuracy',0)*100:.1f}%")

    doc.add_heading("Model 3: Failure Prediction", level=2)
    doc.add_paragraph(f"Accuracy: {fp.get('accuracy',0)*100:.1f}% | AUC-ROC: {fp.get('auc_roc',0):.3f} | AUC-PR: {fp.get('auc_pr',0):.3f}")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 7. METHODOLOGY
    # ════════════════════════════════════════════════════════
    doc.add_heading("7. Methodology & Standards", level=1)

    doc.add_heading("Standards Applied", level=2)
    standards = [
        "IEEE C57.104-2019 \u2014 Gas concentration thresholds (Status 1/2/3), rate limits, Rogers ratios, Key Gas method",
        "IEC 60599:2022 \u2014 Fault taxonomy (PD/D1/D2/T1/T2/T3/DT), Duval Triangle, IEC ratio method",
        "CIGRE TB 771 (2019) \u2014 Duval Pentagon 5-gas centroid classification",
        "IEC 60422:2013 \u2014 Oil quality thresholds (BDV, moisture, acidity, tan delta)",
    ]
    for s in standards:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Health Index Formula", level=2)
    doc.add_paragraph("CHI = 50% \u00D7 DGAF + 15% \u00D7 BDV + 15% \u00D7 Moisture + 10% \u00D7 Acidity + 10% \u00D7 Tan Delta")

    doc.add_heading("Data Integrity", level=2)
    integrity = [
        "All gas values from actual last sample date per transformer (not aggregated)",
        "Fault labels generated by 5-method weighted consensus voting",
        "Temporal train/test split for failure prediction (no data leakage)",
        "Gas data completeness varies: Key Gas 83%, Duval methods 10\u201315%",
        "Oil BDV data sparse (1.5%) \u2014 CHI redistributes weight proportionally",
    ]
    for item in integrity:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Research References", level=2)
    refs = [
        "[1] ML-based multi-method DGA interpretation \u2014 PMC/ScienceDirect, 2024",
        "[2] Fault Classification via DGA and ML: Systematic Review \u2014 MDPI Applied Sciences, 2025",
        "[3] Evaluation of Duval Pentagon \u2014 Springer Electrical Engineering, 2025",
        "[4] Review of Transformer Health Index \u2014 MDPI Electronics, 2023",
        "[5] Modified DGA Scoring with Rate Values \u2014 MDPI Energies, 2024",
        "[6] SHAP + LGBM for DGA Fault Diagnosis \u2014 Energy Informatics, 2025",
        "[7] SMOTE + GBDT Hybrid for DGA \u2014 Arabian J. Sci. Eng., 2025",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"\u2014 End of Report \u2014")
    r.italic = True
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated: {TODAY} | AP Transco DGA Dashboard \u2014 Phase A Local Analysis")
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY

    # Save
    out = os.path.join(REPORTS_DIR, "AP_Transco_DGA_Fleet_Health_Report.docx")
    doc.save(out)
    size = os.path.getsize(out) / 1024
    print(f"  Saved: {out} ({size:.0f} KB)")


if __name__ == "__main__":
    print("Generating Word Report...")
    main()
    print("Done!")
